import streamlit as st
import requests
from bs4 import BeautifulSoup
from duckduckgo_search import DDGS
from urllib.parse import urlparse
from difflib import SequenceMatcher
import io
import re
from collections import defaultdict

# Impor modul AI eksternal (gemini.py)
from gemini import penulis_profesional_ai, jawab_pertanyaan_chat, dapatkan_sinonim_ai

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document

# ==========================================
# 1. SEMANTIC RETRIEVAL & SOURCE INTELLIGENCE
# ==========================================
@st.cache_data(ttl=3600, show_spinner=False)
def expand_semantic_keywords(kata_kunci):
    """Dynamic AI Semantic Expander untuk memperluas kueri lintas bahasa & sinonim."""
    kunci_lower = kata_kunci.lower()
    expanded_set = set([kunci_lower])
    
    for word in kunci_lower.split():
        if len(word) > 2:
            expanded_set.add(word)
            
    sinonim_dari_ai = dapatkan_sinonim_ai(kata_kunci)
    for syn in sinonim_dari_ai:
        if len(syn) > 2:
            expanded_set.add(syn)
            
    return list(expanded_set)

def hitung_source_intelligence_score(url, soup_obj, teks_artikel):
    """Source Intelligence Engine: Menghitung Trust Score 0-99 berbasis multi-parameter."""
    parsed_url = urlparse(url)
    domain = parsed_url.netloc.lower()
    skor_dasar = 50
    if parsed_url.scheme == "https":
        skor_dasar += 10
    if ".go.id" in domain:
        skor_dasar += 25
    elif ".ac.id" in domain or ".edu" in domain:
        skor_dasar += 23
    elif "wikipedia.org" in domain:
        skor_dasar += 20
    elif any(d in domain for d in ["kompas.com", "detik.com", "tempo.co", "antaranews.com", "liputan6.com", "cnnindonesia.com", "katadata.co.id", "bbc.com", "reuters.com", "github.com", "stackoverflow.com", "medium.com"]):
        skor_dasar += 20
    elif any(d in domain for d in ["blogspot.com", "wordpress.com"]):
        skor_dasar -= 15
    else:
        skor_dasar += 10
        
    if len(teks_artikel) > 1000:
        skor_dasar += 10
    elif len(teks_artikel) > 500:
        skor_dasar += 5
    else:
        skor_dasar -= 5
        
    trust_score = max(10, min(skor_dasar, 99))
    if trust_score >= 85:
        label_kualitas = f"🟢 Sangat Kredibel (Trust Score: {trust_score})"
    elif trust_score >= 70:
        label_kualitas = f"🔵 Kredibel & Terverifikasi (Trust Score: {trust_score})"
    elif trust_score >= 50:
        label_kualitas = f"🟡 Cukup / Standar (Trust Score: {trust_score})"
    else:
        label_kualitas = f"🔴 Rendah / Blog (Trust Score: {trust_score})"
    return trust_score, label_kualitas

def cari_sumber_mentah(kueri_asli, daftar_ekspansi):
    """
    PERBAIKAN BUG 2: Multi-Query Search Engine dengan Regional Fallback (id-id -> wt-wt).
    Mencari secara lokal (`id-id`) terlebih dahulu, jika hasil kurang/kosong, 
    otomatis fallback atau menggabungkannya dengan pencarian global (`wt-wt`).
    """
    try:
        links_unik = {}
        kueri_pencarian = [kueri_asli] + daftar_ekspansi[:3]
        kueri_pencarian = list(dict.fromkeys(kueri_pencarian))
        
        with DDGS() as ddgs:
            for kueri in kueri_pencarian:
                # 1. Coba region Indonesia (id-id)
                results = list(ddgs.text(kueri, region="id-id", max_results=5))
                
                # 2. Jika hasil lokal terlalu sedikit (biasanya topik internasional seperti Rust/TensorFlow),
                # fallback/gabungkan dengan pencarian global (wt-wt)
                if len(results) < 3:
                    results_global = list(ddgs.text(kueri, region="wt-wt", max_results=7))
                    results = results + results_global
                    
                for r in results:
                    url = r.get('href')
                    if url and url not in links_unik:
                        links_unik[url] = {
                            'url': url, 
                            'title': r.get('title', 'Sumber Web'), 
                            'snippet': r.get('body', '')
                        }
        return list(links_unik.values())
    except Exception as e:
        st.error(f"Terjadi kesalahan saat mencari API DuckDuckGo: {e}")
        return []

def hitung_kemiripan(teks1, teks2):
    return SequenceMatcher(None, teks1.lower(), teks2.lower()).ratio()

def filter_relevansi_dan_duplikat(sumber_mentah, kata_kunci):
    """
    PERBAIKAN BUG 1: Sistem Skor Kecocokan & Perankingan Berdasarkan Jumlah Keyword.
    Menghitung skor match setiap sumber, menyaring yang memiliki match >= 1, 
    dan meranking ulang sumber dari skor tertinggi ke terendah.
    """
    expanded_keywords = expand_semantic_keywords(kata_kunci)
    scored_sources = []
    
    for item in sumber_mentah:
        teks_gabungan = (item['title'] + " " + item['snippet']).lower()
        
        match_score = 0
        for kw in expanded_keywords:
            if kw in teks_gabungan:
                match_score += 1
                
        if match_score >= 1 or not expanded_keywords:
            scored_sources.append({
                'item': item,
                'score': match_score
            })
            
    scored_sources = sorted(scored_sources, key=lambda x: x['score'], reverse=True)
    
    sumber_terfilter = []
    for scored in scored_sources:
        item = scored['item']
        duplikat = any(hitung_kemiripan(item['title'], unik['title']) >= 0.6 for unik in sumber_terfilter)
        if not duplikat:
            sumber_terfilter.append(item)
            
    return sumber_terfilter[:9]

# ==========================================
# 2. PYTHON RESEARCH ENGINES (ADVANCED)
# ==========================================
def pecah_menjadi_atomic_claims(kalimat):
    titik_pecah = re.split(r',|\bsetelah\b|\bsebelum\b|\bdimana\b|\bketika\b|\bserta\b|\bdan juga\b', kalimat)
    atomic_list = []
    for bagian in titik_pecah:
        bersih = bagian.strip(" .")
        if len(bersih) > 12:
            atomic_list.append(bersih[0].upper() + bersih[1:])
    return atomic_list if atomic_list else [kalimat.strip(".")]

def parsing_causal_relations(kalimat):
    """Causal Engine: Mendeteksi rantai sebab akibat secara presisi."""
    causal_chains = []
    kalimat_lower = kalimat.lower()
    causal_markers = [
        "menyebabkan", "mengakibatkan", "berdampak pada", "karena", 
        "akibat", "memicu", "berkontribusi pada", "sehingga", 
        "oleh sebab itu", "berujung pada", "berimbas pada", 
        "berimbas ke", "mempercepat", "mendorong", "menghasilkan"
    ]
    for marker in causal_markers:
        pola_marker = f" {marker} "
        if pola_marker in kalimat_lower:
            parts = kalimat_lower.split(pola_marker)
            if len(parts) == 2:
                sebab = parts[0].strip(" ,.").title()
                akibat = parts[1].strip(" ,.").title()
                if 3 < len(sebab) < 50 and 3 < len(akibat) < 50:
                    causal_chains.append((sebab, marker, akibat))
                    break
    return causal_chains

def parsing_temporal_lanjutan(kalimat):
    events = []
    kalimat_lower = kalimat.lower()
    for y in re.findall(r'\b(1[0-9]{3}|20[0-9]{2})\b', kalimat):
        events.append((int(y), f"Tahun {y}: {kalimat.strip()}"))
    for a in re.findall(r'abad\s*(?:ke-)?([0-9ivxlc]+)', kalimat_lower):
        num_val = int(a) if a.isdigit() else {'iv': 4, 'v': 5, 'vi': 6, 'vii': 7, 'viii': 8, 'ix': 9, 'x': 10, 'xi': 11, 'xii': 12, 'xiii': 13, 'xiv': 14, 'xv': 15, 'xvi': 16, 'xvii': 17, 'xviii': 18, 'xix': 19, 'xx': 20}.get(a, 5)
        events.append(((num_val - 1) * 100, f"Abad ke-{num_val}: {kalimat.strip()}"))
    return events

def ekstrak_knowledge_graph_python(teks_full):
    """Knowledge Graph Engine: Mengekstrak relasi entitas secara komprehensif."""
    triples = []
    relasi_verbs = [
        "memimpin", "didirikan oleh", "dibangun oleh", "terletak di", "berpusat di", 
        "dipimpin oleh", "mencapai", "menaklukkan", "berdiri pada", "merupakan", "adalah",
        "menjadi", "berasal dari", "berafiliasi dengan", "dipengaruhi oleh", "dipilih sebagai", 
        "ditetapkan sebagai", "diumumkan pada", "diangkat menjadi", "diangkat sebagai",
        "dikenal sebagai", "berkaitan dengan", "berbatasan dengan", "diciptakan oleh", 
        "ditulis oleh", "ditemukan oleh", "menggantikan", "digantikan oleh"
    ]
    for kal in [k.strip() for k in teks_full.split('.') if len(k.strip()) > 20]:
        kal_lower = kal.lower()
        for verb in relasi_verbs:
            pola_verb = f" {verb} "
            if pola_verb in kal_lower:
                parts = kal_lower.split(pola_verb)
                if len(parts) == 2:
                    subjek = parts[0].strip().title()
                    objek = parts[1].strip(". ").title()
                    if 2 < len(subjek) < 40 and 2 < len(objek) < 40:
                        triples.append((subjek, verb, objek))
                        break
    return list(set(triples))[:15]

def ekstrak_entitas_python(teks_full):
    tokoh, lokasi, organisasi, tanggal = set(), set(), set(), set()
    for y in re.findall(r'\b(1[0-9]{3}|20[0-9]{2})\b', teks_full):
        tanggal.add(str(y))
    for w in re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', teks_full):
        if w not in {"Dan", "Yang", "Dari", "Dalam", "Pada", "Dengan", "Untuk", "Sebagai", "Oleh", "Adalah", "Namun", "Selain", "Ketika", "Setelah", "Sebelum", "Karena", "Sehingga"} and len(w) > 3:
            w_lower = w.lower()
            if any(org in w_lower for org in ["kerajaan", "republik", "pt", "cv", "universitas", "badan"]):
                organisasi.add(w)
            elif any(loc in w_lower for loc in ["kota", "kabupaten", "provinsi", "pulau", "candi", "majapahit", "borobudur", "indonesia"]):
                lokasi.add(w)
            else:
                tokoh.add(w)
    return {"Tokoh": list(tokoh)[:6], "Lokasi": list(lokasi)[:6], "Organisasi": list(organisasi)[:6], "Tanggal": sorted(list(tanggal))}

def proses_peneliti_python(url, kata_kunci, source_id):
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        response = requests.get(url, headers=headers, timeout=7)
        soup = BeautifulSoup(response.text, 'html.parser')
        for elemen in soup(['script', 'style', 'header', 'footer', 'nav', 'aside', 'form']):
            elemen.decompose()
            
        paragraf = [p.get_text().strip() for p in soup.find_all('p') if len(p.get_text().strip()) > 30]
        if not paragraf:
            return "", [], [], {}, [], [], 50, "Cukup"
            
        teks_full = " ".join(paragraf)
        trust_score, label = hitung_source_intelligence_score(url, soup, teks_full)
        entities = ekstrak_entitas_python(teks_full)
        triples = ekstrak_knowledge_graph_python(teks_full)
        
        expanded_kw = expand_semantic_keywords(kata_kunci)
        claims, events, causal_list = [], [], []
        
        for p in paragraf:
            for kal in [k.strip() for k in p.split('.') if len(k.strip()) > 20]:
                kal_lower = kal.lower()
                ada_keyword = any(kw in kal_lower for kw in expanded_kw)
                ada_angka = bool(re.search(r'\b(1[0-9]{3}|20[0-9]{2})\b', kal))
                
                if ada_keyword or ada_angka:
                    for ac in pecah_menjadi_atomic_claims(kal):
                        if ac not in claims:
                            claims.append(ac)
                            
                events.extend(parsing_temporal_lanjutan(kal))
                causal_list.extend(parsing_causal_relations(kal))
                
        if not claims and paragraf:
            for p in paragraf[:5]:
                for ac in pecah_menjadi_atomic_claims(p):
                    if ac not in claims:
                        claims.append(ac)
                        
        return ". ".join(claims[:12]) + ".", claims[:12], sorted(list(set(events)), key=lambda x: x[0]), entities, triples, list(set(causal_list)), trust_score, label
    except Exception:
        return "", [], [], {}, [], [], 50, "Cukup"

def build_question_planner_blueprint(query, expanded_kw, sumber_data_lengkap):
    """Dynamic Question Planner Engine: Menyesuaikan sub-pertanyaan berdasarkan kategori topik."""
    gabungan_kata = " ".join(expanded_kw).lower() + " " + query.lower()
    
    if any(k in gabungan_kata for k in ["python", "docker", "tensorflow", "software", "ai", "teknologi", "aplikasi", "programming", "cloud", "linux", "kode", "hugging face", "rust", "sistem", "wicca"]):
        sub_questions = [
            {"domain": "Arsitektur & Teknis", "keyword": ["arsitektur", "teknis", "sistem", "kode", "perangkat", "infrastruktur", "komponen", "cara kerja"]},
            {"domain": "Fitur & Kapabilitas", "keyword": ["fitur", "fungsi", "kemampuan", "mendukung", "kinerja", "keunggulan", "performa"]},
            {"domain": "Ekosistem & Integrasi", "keyword": ["ekosistem", "integrasi", "library", "komunitas", "dukungan", "kompatibilitas"]},
            {"domain": "Implementasi & Use Case", "keyword": ["implementasi", "penggunaan", "solusi", "dipakai", "industri", "contoh"]}
        ]
    elif any(k in gabungan_kata for k in ["kesehatan", "medis", "penyakit", "obat", "virus", "klinis", "terapi", "gejala", "sindrom"]):
        sub_questions = [
            {"domain": "Gejala & Penyebab", "keyword": ["gejala", "penyebab", "virus", "bakteri", "faktor", "risiko", "memicu"]},
            {"domain": "Diagnosis & Penanganan", "keyword": ["diagnosis", "penanganan", "terapi", "obat", "perawatan", "tindakan"]},
            {"domain": "Pencegahan & Mitigasi", "keyword": ["pencegahan", "vaksin", "mitigasi", "menghindari", "protokol"]},
            {"domain": "Dampak Jangka Panjang", "keyword": ["dampak", "komplikasi", "efek", "pemulihan", "jangka panjang"]}
        ]
    elif any(k in gabungan_kata for k in ["sejarah", "kerajaan", "perang", "runtuh", "pemberontakan", "kemerdekaan", "presiden", "politik", "negara"]):
        sub_questions = [
            {"domain": "Politik & Kekuasaan", "keyword": ["politik", "raja", "tahta", "pemerintahan", "kekuasaan", "konflik", "pemberontakan"]},
            {"domain": "Ekonomi & Sumber Daya", "keyword": ["ekonomi", "pajak", "perdagangan", "rempah", "pelabuhan", "upeti"]},
            {"domain": "Militer & Pertahanan", "keyword": ["militer", "perang", "pasukan", "senjata", "benteng", "ekspansi"]},
            {"domain": "Sosial & Budaya", "keyword": ["agama", "sosial", "budaya", "masyarakat", "candi", "tradisi"]}
        ]
    else:
        sub_questions = [
            {"domain": "Latar Belakang & Konteks", "keyword": ["awal", "munculnya", "latar belakang", "dasar", "asal", "definisi"]},
            {"domain": "Faktor Pendorong (Sebab)", "keyword": ["karena", "menyebabkan", "mendorong", "alasan", "faktor", "pemicu"]},
            {"domain": "Proses & Mekanisme (Cara)", "keyword": ["proses", "cara", "sistem", "metode", "berlangsung", "berjalan"]},
            {"domain": "Dampak & Konsekuensi (Akibat)", "keyword": ["dampak", "akibat", "hasil", "pengaruh", "perubahan", "implikasi"]}
        ]

    blueprint = []
    for sq in sub_questions:
        matched_claims = []
        for item in sumber_data_lengkap:
            for claim in item['claims']:
                if any(kw in claim.lower() for kw in sq['keyword']):
                    matched_claims.append((item['id'], claim))
        blueprint.append({'domain': sq['domain'], 'matched_claims': matched_claims[:4]})
        
    return blueprint

def hitung_statistik_riset(sumber_data):
    """Menghitung Confidence & Konsensus riset dengan Safety Net anti-nol."""
    jumlah_sumber = len(sumber_data)
    if jumlah_sumber == 0:
        return 10, 0.0, 0, 0
        
    rata_trust = sum(i.get('trust_score', 50) for i in sumber_data) / jumlah_sumber
    confidence = int(round((rata_trust * 0.4) + (min(jumlah_sumber / 5.0, 1.0) * 30) + 30))
    consensus = int(round(min(rata_trust * 0.98, 98)))
    return confidence, rata_trust, jumlah_sumber, consensus

# ==========================================
# 3. EKSPOR DOKUMEN & TAMPILAN STREAMLIT
# ==========================================
def generate_pdf(topik, skor, consensus, rata_skor, total_sumber, hasil_analisis, sumber_list):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    story = [Paragraph(f"LAPORAN RISET: {topik.upper()}", ParagraphStyle('T', fontSize=16, textColor=colors.HexColor("#1E3A8A"))), Spacer(1, 10)]
    for para in hasil_analisis.split("\n"):
        if para.strip():
            story.append(Paragraph(para, ParagraphStyle('N', fontSize=10, leading=14)))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_docx(topik, skor, consensus, rata_skor, total_sumber, hasil_analisis, sumber_list):
    doc = Document()
    doc.add_heading(f"Laporan Riset: {topik}", level=1)
    doc.add_paragraph(hasil_analisis)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

st.set_page_config(page_title="AI Researcher - Enterprise Edition", page_icon="⚡")
st.title("⚡ Enterprise AI Researcher Engine")
st.caption("Sistem Riset Mandiri Berbasis Multi-Query Semantic, Question Planner, & Causal Analyzer")

query = st.text_input("Masukkan topik riset (contoh: Wicca, Rust, Hugging Face, PPPK paruh waktu, dll):")

if st.button("Jalankan Riset Lengkap"):
    if query.strip():
        with st.status("Menjalankan Peneliti Python & Gemini AI...", expanded=True) as status:
            
            st.write("🧠 AI Merumuskan Sinonim Semantik (Semantic Expansion)...")
            kata_kunci_ekspansi = expand_semantic_keywords(query)
            st.write(f"*Padanan yang ditemukan AI:* `{', '.join(kata_kunci_ekspansi)}`")
            
            st.write("🔍 Menjalankan Multi-Query Search (dengan Regional Fallback)...")
            sumber_mentah = cari_sumber_mentah(query, kata_kunci_ekspansi)
            if not sumber_mentah:
                status.update(label="Tidak ada sumber ditemukan. Periksa koneksi internet.", state="error")
                st.stop()
                
            st.write("🧹 Menyaring duplikat & meranking sumber berdasarkan Skor Keyword...")
            sumber_list = filter_relevansi_dan_duplikat(sumber_mentah, query)
            if not sumber_list:
                status.update(label="Tidak ada sumber yang relevan dengan kueri tersebut.", state="error")
                st.stop()
                
            sumber_data_lengkap = []
            timeline_global = []
            causal_chains_global = []
            global_entities = {"Tokoh": [], "Lokasi": [], "Organisasi": [], "Tanggal": []}
            global_triples = []
            
            for idx, item in enumerate(sumber_list, 1):
                res = proses_peneliti_python(item['url'], query, idx)
                if res[0] or res[1] or res[3]:
                    sumber_data_lengkap.append({
                        'id': idx, 'title': item['title'], 'url': item['url'],
                        'trust_score': res[6], 'kualitas': res[7], 'claims': res[1] if res[1] else [item['snippet']]
                    })
                    timeline_global.extend(res[2])
                    global_triples.extend(res[4])
                    causal_chains_global.extend(res[5])
                    for k in global_entities:
                        global_entities[k].extend(res[3].get(k, []))
            
            if not sumber_data_lengkap and sumber_list:
                for idx, item in enumerate(sumber_list, 1):
                    sumber_data_lengkap.append({
                        'id': idx, 'title': item['title'], 'url': item['url'],
                        'trust_score': 60, 'kualitas': "🟡 Standar", 'claims': [item['snippet']]
                    })
            
            for k in global_entities:
                global_entities[k] = sorted(list(set(global_entities[k])))[:8]
            global_triples = list(set(global_triples))
            causal_chains_global = list(set(causal_chains_global))
            
            st.write("🧭 Menyusun Question Planner Blueprint...")
            planner_blueprint = build_question_planner_blueprint(query, kata_kunci_ekspansi, sumber_data_lengkap)
            timeline_global = sorted(list(set(timeline_global)), key=lambda x: x[0])
            
            score_akhir, rata_skor_trust, total_sumber, consensus_score = hitung_statistik_riset(sumber_data_lengkap)
            
            st.write("✍️ Memanggil modul AI Penulis (gemini.py) untuk menyusun laporan...")
            hasil_analisis, model_pakai = penulis_profesional_ai(
                planner_blueprint, causal_chains_global, timeline_global, global_entities, global_triples, query, score_akhir
            )
            
            if model_pakai:
                st.write(f"✨ Berhasil menggunakan model: `{model_pakai}`")
            
            status.update(label="Riset Berhasil Tuntas!", state="complete", expanded=False)
            
        st.session_state['hasil_analisis'] = hasil_analisis
        st.session_state['sumber_data_lengkap'] = sumber_data_lengkap
        st.session_state['planner_blueprint'] = planner_blueprint
        st.session_state['causal_chains_global'] = causal_chains_global
        st.session_state['global_entities'] = global_entities
        st.session_state['timeline_global'] = timeline_global
        st.session_state['global_triples'] = global_triples
        st.session_state['score_akhir'] = score_akhir
        st.session_state['rata_skor_trust'] = rata_skor_trust
        st.session_state['total_sumber'] = total_sumber
        st.session_state['consensus_score'] = consensus_score
        st.session_state['query'] = query
        st.session_state['kata_kunci_ekspansi'] = kata_kunci_ekspansi
        st.session_state['messages'] = [{"role": "assistant", "content": f"Laporan untuk **{query}** berhasil disusun."}]

if 'hasil_analisis' in st.session_state:
    st.markdown("---")
    st.subheader("📈 Dashboard Statistik & Kesehatan Riset")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Sumber Unik", st.session_state['total_sumber'])
    with col2:
        st.metric("Avg Trust Score", f"{st.session_state['rata_skor_trust']:.1f}/100")
    with col3:
        st.metric("Konsensus", f"{st.session_state['consensus_score']}%")
    with col4:
        st.metric("Confidence", f"{st.session_state['score_akhir']}%")
    st.progress(st.session_state['score_akhir'] / 100.0)

    st.markdown("---")
    st.subheader("📥 Unduh Laporan")
    col_pdf, col_docx = st.columns(2)
    pdf_bytes = generate_pdf(st.session_state['query'], st.session_state['score_akhir'], st.session_state['consensus_score'], st.session_state['rata_skor_trust'], st.session_state['total_sumber'], st.session_state['hasil_analisis'], st.session_state['sumber_data_lengkap'])
    docx_bytes = generate_docx(st.session_state['query'], st.session_state['score_akhir'], st.session_state['consensus_score'], st.session_state['rata_skor_trust'], st.session_state['total_sumber'], st.session_state['hasil_analisis'], st.session_state['sumber_data_lengkap'])
    
    with col_pdf:
        st.download_button("📄 Unduh PDF", data=pdf_bytes, file_name=f"Riset_{st.session_state['query'].replace(' ', '_')}.pdf", mime="application/pdf")
    with col_docx:
        st.download_button("📝 Unduh Word", data=docx_bytes, file_name=f"Riset_{st.session_state['query'].replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.subheader(f"📊 Laporan Riset: {st.session_state['query']}")
    st.markdown(st.session_state['hasil_analisis'])

    with st.expander("🧠 Lihat Ekspansi Semantik AI (Sinonim yang Ditemukan)"):
        st.markdown(f"**Query Asli:** `{st.session_state['query']}`")
        st.markdown(f"**Ditangkap AI sebagai:** `{', '.join(st.session_state['kata_kunci_ekspansi'])}`")

    st.markdown("---")
    st.subheader("💬 Tanya Jawab Berbasis Perencanaan Agen")
    for message in st.session_state.get("messages", []):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if user_prompt := st.chat_input("Tanyakan detail sub-dimensi atau analisis riset..."):
        st.session_state["messages"].append({"role": "user", "content": user_prompt})
        with st.chat_message("user"):
            st.markdown(user_prompt)
        with st.chat_message("assistant"):
            with st.spinner("Menyusun jawaban agen..."):
                jawaban = jawab_pertanyaan_chat(
                    user_prompt, 
                    st.session_state['planner_blueprint'], 
                    st.session_state['causal_chains_global'], 
                    st.session_state['timeline_global'], 
                    st.session_state['hasil_analisis']
                )
                st.markdown(jawaban)
                st.session_state["messages"].append({"role": "assistant", "content": jawaban})
